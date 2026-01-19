import pandas as pd
import numpy as np
from django.shortcuts import render
from django.http import HttpResponse
from docx import Document
from .forms import UploadForm


# ---------------- UPLOAD PAGE ----------------
def upload_data(request):
    form = UploadForm()
    return render(request, "upload_data.html", {"form": form})


# ---------------- ANALYZE EXCEL ----------------
def analyze(request):
    if request.method == "POST":
        form = UploadForm(request.POST, request.FILES)
        if not form.is_valid():
            return render(request, "upload_data.html", {"form": form})

        df = pd.read_excel(request.FILES["file"])
        total_n = len(df)

        raw_cols = df.columns.tolist()
        raw_data = df.values.tolist()

        numeric = df.select_dtypes(include=np.number)
        categorical = df.select_dtypes(exclude=np.number)

        table_discussions = []

        # ---------- STATISTICS (DERIVED) ----------
        stats_table = []
        stats_disc = []

        if not numeric.empty:
            stats = numeric.agg(["mean", "std"]).T.round(2)
            for v, row in stats.iterrows():
                stats_table.append([v, row["mean"], row["std"]])
                stats_disc.append(
                    f"The mean value of {v} is {row['mean']} with a standard deviation of {row['std']}, "
                    f"indicating the central tendency and variability of the respondents."
                )

        table_discussions.append(("Table 2. Statistical Summary (Derived)", stats_disc))

        # ---------- FREQUENCY (SUMMARY) ----------
        freq_tables = []
        for col in categorical.columns:
            freq = df[col].value_counts(dropna=False).reset_index()
            freq.columns = ["Category", "Frequency"]
            freq["Percent"] = round(freq["Frequency"] / total_n * 100, 2)

            major = freq.iloc[0]
            desc = (
                f"The distribution of {col} shows that the largest group is {major['Category']} "
                f"with {major['Frequency']} respondents ({major['Percent']}%), "
                f"indicating that this category is the most prevalent among the participants."
            )

            freq_tables.append((col, freq.values.tolist(), desc))
            table_discussions.append((f"Frequency Distribution: {col}", [desc]))

        # ---------- CROSSTAB (CLASSIFICATION) ----------
        cross_tables = []
        cat_cols = categorical.columns.tolist()

        for i in range(len(cat_cols)):
            for j in range(i + 1, len(cat_cols)):
                ct = pd.crosstab(df[cat_cols[i]], df[cat_cols[j]])
                rows = [(r, ct.loc[r].tolist()) for r in ct.index]

                desc = (
                    f"This cross-tabulation presents the relationship between {cat_cols[i]} and {cat_cols[j]}, "
                    f"showing how the distribution of one variable differs across the categories of the other."
                )

                cross_tables.append((f"{cat_cols[i]} × {cat_cols[j]}", ct.columns.tolist(), rows, desc))
                table_discussions.append((f"Cross Tabulation: {cat_cols[i]} × {cat_cols[j]}", [desc]))

        # ---------- CORRELATION ----------
        corr_cols, corr_rows = None, None
        corr_disc = []

        if numeric.shape[1] > 1:
            corr = numeric.corr().round(3)
            corr_cols = corr.columns.tolist()
            corr_rows = [(i, r.tolist()) for i, r in corr.iterrows()]

            for rname, vals in corr_rows:
                for cname, v in zip(corr_cols, vals):
                    if rname != cname:
                        strength = "weak"
                        if abs(v) >= 0.7:
                            strength = "strong"
                        elif abs(v) >= 0.4:
                            strength = "moderate"

                        corr_disc.append(
                            f"There is a {strength} relationship between {rname} and {cname} (r = {v})."
                        )

            table_discussions.append(("Correlation Matrix", corr_disc))

        # SAVE SESSION
        request.session["report_data"] = {
            "title": form.cleaned_data["title"],
            "objective": form.cleaned_data["objective"],
            "problem": form.cleaned_data["problem"],
            "raw_cols": raw_cols,
            "raw_data": raw_data,
            "stats_table": stats_table,
            "freq_tables": freq_tables,
            "cross_tables": cross_tables,
            "corr_cols": corr_cols,
            "corr_rows": corr_rows,
            "table_discussions": table_discussions,
            "total_n": total_n
        }

        return render(request, "results.html", {
            "title": form.cleaned_data["title"],
            "objective": form.cleaned_data["objective"],
            "problem": form.cleaned_data["problem"],
            "raw_cols": raw_cols,
            "raw_data": raw_data,
            "stats_table": stats_table,
            "freq_tables": freq_tables,
            "cross_tables": cross_tables,
            "corr_cols": corr_cols,
            "corr_rows": corr_rows,
            "table_discussions": table_discussions,
            "total_n": total_n
        })


# ---------------- EXPORT TO WORD ----------------
def export_word(request):
    data = request.session.get("report_data")
    if not data:
        return HttpResponse("No data available.")

    doc = Document()
    doc.add_heading(data["title"], level=1)

    doc.add_paragraph(f"Objective: {data['objective']}")
    doc.add_paragraph(f"Problem Statement: {data['problem']}")
    doc.add_paragraph(f"Number of Respondents: {data['total_n']}")

    # RAW DATA
    doc.add_heading("Table 1. Raw Data (Original)", level=2)
    t = doc.add_table(rows=1, cols=len(data["raw_cols"]))
    for i, c in enumerate(data["raw_cols"]):
        t.rows[0].cells[i].text = str(c)
    for r in data["raw_data"]:
        row = t.add_row().cells
        for i, v in enumerate(r):
            row[i].text = str(v)

    # STATISTICS
    doc.add_heading("Table 2. Statistical Summary (Derived)", level=2)
    t = doc.add_table(rows=1, cols=3)
    t.rows[0].cells[0].text = "Variable"
    t.rows[0].cells[1].text = "Mean"
    t.rows[0].cells[2].text = "Std Dev"
    for r in data["stats_table"]:
        row = t.add_row().cells
        row[0].text = str(r[0])
        row[1].text = str(r[1])
        row[2].text = str(r[2])

    # FREQUENCY
    for col, rows, desc in data["freq_tables"]:
        doc.add_heading(f"Frequency Distribution: {col}", level=2)
        doc.add_paragraph(desc)
        t = doc.add_table(rows=1, cols=3)
        t.rows[0].cells[0].text = "Category"
        t.rows[0].cells[1].text = "Frequency"
        t.rows[0].cells[2].text = "Percent"
        for r in rows:
            row = t.add_row().cells
            row[0].text = str(r[0])
            row[1].text = str(r[1])
            row[2].text = str(r[2])

    # CROSSTAB
    for title, cols, rows, desc in data["cross_tables"]:
        doc.add_heading(f"Cross Tabulation: {title}", level=2)
        doc.add_paragraph(desc)
        t = doc.add_table(rows=1, cols=len(cols) + 1)
        t.rows[0].cells[0].text = ""
        for i, c in enumerate(cols):
            t.rows[0].cells[i + 1].text = str(c)
        for rname, vals in rows:
            row = t.add_row().cells
            row[0].text = str(rname)
            for i, v in enumerate(vals):
                row[i + 1].text = str(v)

    # CORRELATION
    if data.get("corr_cols"):
        doc.add_heading("Correlation Matrix", level=2)
        t = doc.add_table(rows=1, cols=len(data["corr_cols"]) + 1)
        t.rows[0].cells[0].text = ""
        for i, c in enumerate(data["corr_cols"]):
            t.rows[0].cells[i + 1].text = str(c)
        for rname, vals in data["corr_rows"]:
            row = t.add_row().cells
            row[0].text = str(rname)
            for i, v in enumerate(vals):
                row[i + 1].text = str(v)

    # DISCUSSIONS
    doc.add_heading("Discussion of Results", level=1)
    for title, texts in data["table_discussions"]:
        doc.add_heading(title, level=2)
        for t in texts:
            doc.add_paragraph(t)

    response = HttpResponse(
        content_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
    )
    response["Content-Disposition"] = "attachment; filename=Research_Report.docx"
    doc.save(response)
    return response
